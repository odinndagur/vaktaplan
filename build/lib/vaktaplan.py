#!/Users/odinndagur/.odb/vaktaplan/camelot-env/bin/python3

# IMPORTS
import camelot
import pandas as pd
import os
import re
from numpy import nan
import pdfplumber
import datetime
import docx
from pdf2image import convert_from_path
from docx.enum.text import WD_COLOR_INDEX

ppl = {
'Agnes Rut Ásgeirsdóttir': '',
'Arndís prufa': 'Arndís',
'Aron Freyr Kristjánsson': 'Aron',
'Björg Hákonardóttir': 'Björg',
'Breki prufa': 'Breki',
'Daníel Pétursson': 'Daníel',
'Davíð Arnljótur Karlsson': 'Davíð',
'Friðberg Reynir Traustason': 'Friðberg',
'Guðbjartur Máni Gíslason': 'G. Máni',
'Guðni Þór Björnsson': 'Guðni',
'Gunnar Kristinn Þórðarson': 'Gunnar',
'Harpa Mjöll Þórsdóttir': 'Harpa',
'Hrafn Ingi Jóhannsson': 'Hrafn',
'Hrafnhildur S Benediktsdóttir': 'Hrafnhildur',
'Ingibjörg Rún Jóhannesdóttir': 'Ingibjörg Rún',
'Íris Hildur Birgisdóttir': 'Íris',
'Ísak Örn Ívarsson': 'Ísak',
'Jón Arnar Björnsson': 'Nonni',
'Jón Egill Eyþórsson': 'Jón Egill',
'Karl Héðinn Kristjánsson': 'Kalli',
'Kolfinna Kristjánsdóttir': 'Kolfinna',
'Kristín Anna Ólafsdóttir': 'Kristín',
'Kristján Ingi Rúnarsson': 'Kristján Ingi',
'Kristján Kristjánsson': 'Kristján K',
'KVK 1': '',
'Lena María Brynjarsdóttir': 'Lena',
'Óðinn Dagur Bjarnason': 'Óðinn',
'Ólafur Árni Jónsson': 'Ólafur',
'Ragnar Ingi Karlsson': 'Ragnar Ingi',
'Ragnheiður Lindudóttir': 'Ragnheiður',
'Rakel Ketilsdóttir': 'Rakel',
'Róbert Steinar Hjálmarsson': 'Róbert',
'Rósa María Friðriksdóttir': 'Rósa',
'Sara Bergmann Valtýsdóttir': 'Sara',
'Sesselja Hansen Daðadóttir': 'Sesselja',
'Sigríður Eyrún Sigurjónsdóttir': 'Sigga',
'Sigrún Erla Lárusdóttir': 'Sigrún',
'Sigurjón Már Markússon': 'Sigurjón',
'Silja Hjaltadóttir': 'Silja',
'Sindri Rafn Guðmundsson': 'Sindri Rafn',
'Sindri Viborg': 'Sindri V',
'Stefanía Katrín prufa': 'Stefanía K',
'Stefanía Smáradóttir': 'Stefanía Smáradóttir',
'Sveinbjörn Skúlason': 'Sveinbjörn',
'Úlfar Gunnarsson': 'Úlfar',
'Valgerður Selma Lúthersdóttir': 'Valgerður',
'Þorvaldur Jóhannesson': 'Þorvaldur',
'Þórir Óskar Björnsson': 'Skari',
'Ævar Örn Erlendsson': 'Ævar',
'övantar': '',
}

verbose = 1

def printv(input_str:str):
    if verbose == 1:
        print(input_str)

def printvv(input_str:str):
    if verbose > 1:
        print(input_str)



def main():
    import argparse
    import sys
    global plan

    parser = argparse.ArgumentParser()
    parser.add_argument('-i','--input',help='Input file to work on, generally ends with .pdf or .csv.')
    parser.add_argument('-o', '--output', help='Where to save output from program. Defaults to current working directory. It creates a folder by the same name as the plan to save all files.')
    parser.add_argument('-d','--generate-dayplans', dest='generate_dayplans', action='store_true', help='Should we generate dayplans for all the days and output to the drive?')
    parser.add_argument('-s','--shifts',help='Get shifts for person?')
    args = parser.parse_args(sys.argv[1:])

    if not args.input:
        return
    if args.input.endswith('.pdf'):
        plan = Vaktaplan.from_pdf(args.input)
    if args.input.endswith('.csv'):
        plan = Vaktaplan.from_csv(args.input)
    file_base_name = os.path.splitext(args.input)[0]
    
    if args.output:
        output_folder = os.path.join(args.output,file_base_name)
    else:
        output_folder = os.path.join(os.getcwd(),file_base_name)
    os.makedirs(output_folder,exist_ok=True)
    
    if args.generate_dayplans: #tempfile ruglið
        make_and_save_zip_file(save_location='/Users/odinndagur/Desktop/',zip_name="delivery.zip",plan=plan)


def make_and_save_zip_file(save_location,zip_name,plan):
    import os
    import shutil
    import tempfile
    import zipfile
    zip_path = os.path.join(save_location, zip_name)

    zip_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zfd:
            for day in plan.get_days():
                day_path = os.path.join(zip_dir,f'{str(day.date)}.docx')
                day.generate_dayplan().save(day_path)
                zfd.write(day_path, f'{str(day.date)}.docx')
        return zip_path
    finally:
        shutil.rmtree(zip_dir)



def get_year(month:int):
    """
    Helper function to get a year for a date since no years are in the plans.
    """
    currentDate = datetime.datetime.today()
    return currentDate.year if month >= currentDate.month else currentDate.year + 1

def get_weekday(year:int = None, month:int = None, day:int = None) -> str:
    days = ["Mánudagur", "Þriðjudagur", "Miðvikudagur", "Fimmtudagur", "Föstudagur", "Laugardagur", "Sunnudagur"]
    if not any([year,month,day]):
        raise ValueError('Need to input date')
    if not year:
        year = get_year(int(month))
    inputDate = datetime.date(year=year,month=int(month),day=int(day))
    return days[inputDate.weekday()]


class Shift:
    def __init__(
            self,
            person:str,
            shift_type:str,
            start_datetime:datetime.datetime,
            end_datetime:datetime.datetime,
        ):
        self.person = person
        self.shift_type = shift_type
        self.start_datetime = start_datetime
        self.end_datetime = end_datetime
        self.shift_hours = f'{self.start_datetime.time().hour}-{self.end_datetime.time().hour}'
        self.shift_date = f'{self.start_datetime.date()}'
        self.shift_group = self.get_shift_group()
    
    def __str__(self):
        return f'{self.person} {self.shift_hours}'

    def __repr__(self):
        return self.__str__()

    def get_shift_group(self):
        if 0 < self.start_datetime.hour < 12:
            if self.end_datetime.hour > 16:
                return 'DV'
            else:
                return 'MV'
        if 12 <= self.start_datetime.hour < 16:
            return 'DV'
        if 16 <= self.start_datetime.hour < 23:
            return 'KV'
        if 23 <= self.start_datetime.hour <= 24:
            return 'NV'
        if 20 < self.end_datetime.hour <= 24:
            return 'KV'

class Work_day:
    def __init__(self,shifts:list[Shift]):
        self.date:datetime.date = shifts[0].start_datetime.date()
        self.dayplan:docx.Document = self.generate_dayplan()
        self.shifts = shifts

        self.LRL = sorted([shift for shift in self.shifts if shift.shift_type == 'LRL'], key= lambda x: x.start_datetime.hour)
        self.PHA = sorted([shift for shift in self.shifts if shift.shift_type == 'PHA'], key= lambda x: x.start_datetime.hour)
        self.AS = sorted([shift for shift in self.shifts if shift.shift_type == 'AS'], key= lambda x: x.start_datetime.hour)
        self.GR = sorted([shift for shift in self.shifts if shift.shift_type == 'GR'], key= lambda x: x.start_datetime.hour)
        self.BEG = sorted([shift for shift in self.shifts if shift.shift_type == 'BEG'], key= lambda x: x.start_datetime.hour)
        self.GH = sorted([shift for shift in self.shifts if shift.shift_type == 'GH'], key= lambda x: x.start_datetime.hour)
        self.UB = sorted([shift for shift in self.shifts if shift.shift_type == 'UB'], key= lambda x: x.start_datetime.hour)

        self.MV = sorted([shift for shift in self.shifts if shift.shift_group == 'MV'], key= lambda x: x.start_datetime.hour)
        self.DV = sorted([shift for shift in self.shifts if shift.shift_group == 'DV'], key= lambda x: x.start_datetime.hour)
        self.KV = sorted([shift for shift in self.shifts if shift.shift_group == 'KV'], key= lambda x: x.start_datetime.hour)
        self.NV = sorted([shift for shift in self.shifts if shift.shift_group == 'NV'], key= lambda x: x.start_datetime.hour)

    def generate_dayplan(self) -> docx.Document:
        """
        The function to make a word document from a day object.
        """
        # doc = docx.Document(os.path.join(os.path.dirname(os.path.abspath(__file__)),'fim_proto.docx'))
        # doc = docx.Document('fim_proto.docx')
        doc = docx.Document(os.path.join(os.path.dirname(__file__),'fim_proto.docx'))
        for col_idx,col in enumerate(doc.tables[1].columns):
            for cell in doc.tables[1].column_cells(col_idx):
                if cell.text.strip() == 'dags':
                    p = cell.paragraphs[0]
                    p.clear()
                    p.add_run(f'{self.date.strftime("%d.%m.%y")}')
                if cell.text.strip() == 'UB':
                    p = cell.paragraphs[0]
                    p.clear()
                    shifts = [str(shift) for shift in self.UB]
                    p.add_run('\n'.join(shifts))

                    UB_RUN = '\n'.join(shifts)
                    printvv(f'UB - added run:\n{UB_RUN}')
                if cell.text.strip() == 'vikudags':
                    p = cell.paragraphs[0]
                    p.clear()
                    if self.date:
                        weekday = get_weekday(month = self.date.month, day = self.date.day)
                        p.add_run(weekday)
                        printvv(f'vikudags - added run:\n{weekday}')
                if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                    temp = cell.text.strip()
                    shift_type, shift_group = cell.text.strip().split(' ')
                    if not self.__dict__.get(shift_type):
                        break
                    p = cell.paragraphs[0]
                    p.clear()
                    shifts = [str(shift) for shift in self.__dict__[shift_type] if shift.shift_group.lower() == shift_group.lower()]
                    p.add_run('\n'.join(shifts))
                    run = '\n'.join(shifts)


        for col_idx,col in enumerate(doc.tables[1].columns):
            for cell in doc.tables[1].column_cells(col_idx):
                if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                    temp = cell.text.strip().split(' ')[0]
                    cell.paragraphs[0].clear()
                    cell.paragraphs[0].add_run(f'{temp} VANTAR')
                    for run in cell.paragraphs[0].runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return doc



class Vaktaplan:
    """
    Basic shift plan class.
    Can be generated from a pdf (first time) or csv (saved from the class, faster for iteration).
    """
    def __init__(self,df:pd.DataFrame):
        self.df = df
        self.people = self.get_people()
        self.days: list[Work_day] = list(self.get_days())
        self.start_date, self.end_date = self.get_date_range()

    def __str__(self):
        return f'Shiftplan from {self.start_date} to {self.end_date}.'

    def to_csv(self,path:str):
        self.df.to_csv(path)
    
    def get_days(self): #yields generator
        """
        Generator function that iterates over the shift plan
        and makes a strange defaultdict of defaultdicts of lists heh.
        Should make a class for a day since it's all so custom.

        Yields objects with all shifts for the day.
        """
        for idx in range(len(self.df.columns)):
            day_shifts = []
            slice = self.df.iloc[:,idx].replace('',nan).dropna()
            shift_date = slice.name.split('\n')[0]
            if not len(slice):
                continue
            for row_idx, shift in enumerate(slice):
                person = self.get_name(slice.index[row_idx])
                if not shift:
                    break
                if shift == 'ORLOF':
                    break
                time, type = shift.split(' ')

                start_time,end_time = time.split('-')
                start_hours,start_minutes = map(int,start_time.split(':'))
                end_hours,end_minutes = map(int,end_time.split(':'))
                day, month = map(int,shift_date.split('.'))
                year = get_year(month=month)

                start_datetime = datetime.datetime(year=year,month=month,day=day,hour=start_hours,minute=start_minutes)
                end_datetime = datetime.datetime(year=year,month=month,day=day,hour=end_hours,minute=end_minutes)

                day_shifts.append(Shift(person=person,shift_type=type,start_datetime=start_datetime,end_datetime=end_datetime))
            yield Work_day(shifts=day_shifts)

    @classmethod
    def pdf_to_df(cls,file: str) -> pd.DataFrame:
        """
        The actual function that parses the pdf to make a dataframe.
        Uses weird helper functions within the class, needs refactoring.
        """
        print('processing pdf')
        global h,w,new_h,new_w,pdfs
        pdfs = convert_from_path(file)
        with pdfplumber.open(file) as pdf:
            page_1 = pdf.pages[0]
        h,w = page_1.height, page_1.width
        new_h,new_w = pdfs[0].height, pdfs[0].width

        tables = camelot.read_pdf(file,pages='1-end',flavor='lattice',line_scale=50,line_tol=1)
        cls.add_shift_text(tables)
        processed_dfs = [cls.process_df(table) for table in tables]
        # concat fyrst
        concatenated_dfs = [pd.concat(processed_dfs[offset:offset+cls.get_num_pages(tables)]) for offset in range(0,tables.n,cls.get_num_pages(tables))]
        # síðan join
        output_df = concatenated_dfs[0]
        for df in concatenated_dfs[1:]:
            output_df = output_df.join(df)
        return output_df
    
    def get_shifts_for_person(self,person: str) -> str:
        """
        Helper function that returns a list of shifts for a specific person
        """
        shifts = self.df.loc[person].replace('',nan).dropna()
        return [f'{date.split(chr(10))[0]} {shift}' for date,shift in zip(shifts.index,shifts.values)]

    def get_people(self) -> list[str]:
        """
        Helper function for getting a list of people in the plan.
        """
        return [person for person in list(self.df.index) if len(person)]


    @staticmethod
    def is_first_page(df: pd.DataFrame) -> bool:
        """
        Janky custom helper function to see if it's the first page of a dayplan.
        The plans we use have the word "Hæfniþáttur" in the first page of each pdf.
        """
        for x in df.iloc[:,0].values:
            if 'Hæf' in x:
                return True
        return False

    @staticmethod
    def get_first_date_cell(df) -> tuple[int, int]:
        """
        Helper function to get our bearings in the dataframe.
        """
        import re
        for row_idx,row in df.iterrows():
            for col_idx, cell in enumerate(row):
                # print(f'row {row_idx}, col {col_idx}, cell: {cell}')
                if re.match('[0-9][0-9]\.[0-9][0-9]', cell):
                    return col_idx,row_idx

    @classmethod
    def get_num_pages(cls,tables):
        """
        Uses the is_first_page() function to find how many pages
        fit together (in case people were added in later pages).
        """
        counts = []
        last_first_page = 0
        for idx,table in enumerate(tables):
            if cls.is_first_page(table.df):
                if idx > 0:
                    counts.append(idx - last_first_page)
                    last_first_page = idx
        if len(set(counts)) == 1:
            return counts[0]
        else:
            return counts

    @classmethod
    def get_color(cls,img,cell):
        """
        Gets the color of a cell to decide what kind of shift it is.
        """
        y = new_h - ((cell.y1 + 3)/h * new_h)
        x = ((cell.x1 + cell.x2)/2)/w * new_w
        return img.getpixel((x,y))

    @classmethod
    def get_colors_from_tables(cls,tables):
        """
        Returns a set of colors so we can work with the
        individual colors that are in the document.
        """
        colors = set()
        for idx, table in enumerate(tables):
            for row in table.cells:
                for cell in row:
                    colors.add(cls.get_color(pdfs[idx],cell))
        return colors

    colors =  {
        (255, 255, 0): 'GH', #gulur
        (198, 198, 198): '',#ljosgrar
        (240, 240, 240): '',#ljosljosgrar
        (255, 128, 255): 'NV',
        (0, 128, 0): 'BEG',#graenn
        (255, 0, 0): 'GR',#raudur
        (129, 129, 129): '',#mediumgrar
        (122, 122, 122): '',#mediumgrar
        (128, 0, 255): 'LRL',#fjolublar
        (255, 255, 255): 'UB',#hvitur
        (80, 138, 160): 'ORLOF',
        (128, 128, 64): '',
        (128, 128, 128): 'PHA',#mediumgrar
        (128, 0, 64): 'AS',
        }

    @classmethod
    def add_shift_text(cls,tables) -> None:
        """
        Adds the color information in the text of the cell
        so we know what kind of shift it is when it's a csv
        with no colors.
        """
        for page_num,table in enumerate(tables):
            df = table.df
            for row_num,row in enumerate(table.cells):
                for col_num,cell in enumerate(row):
                    cell_color = cls.get_color(pdfs[page_num],cell=cell)
                    if cell_color in cls.colors and (cell.text != '') and re.match('[0-9][0-9]:[0-9][0-9]-[0-9][0-9]:[0-9][0-9]',cell.text):
                        df.iloc[row_num,col_num] += ' ' + cls.colors[cell_color]

    @classmethod
    def process_df(cls,table) -> pd.DataFrame:
        """
        Fix all dataframes so they begin on the correct cell and align
        so they make sense when we join them together later.
        """
        df = table.df.copy()
        x,y = cls.get_first_date_cell(df)
        df.iloc[y,x-1] = 'Starfsmaður'
        df.columns = df.iloc[y,:]
        df = df.iloc[y+1:,x-1:]
        df = df.set_index('Starfsmaður')
        df = df.replace('',nan).dropna(how='all').dropna(how='all',axis=1).replace(nan,'')
        return df

    @classmethod     
    def get_name(cls,person: str):
        """
        So we can use names that aren't the official full name
        (nicknames or first names for instance) in the dayplans.
        """
        # from ppl import ppl
        if person in ppl:
            if ppl[person]:
                return ppl[person]
        if len(person.split(' ')) > 2:
            return ' '.join(person.split(' ')[:2])
        return person

    @classmethod
    def from_pdf(cls,file:str):
        """
        Helper function to generate a Vaktaplan object from a pdf file.
        """
        df = cls.pdf_to_df(file)
        return Vaktaplan(df=df)


    @staticmethod
    def from_csv(file:str):
        """
        Helper function to generate a Vaktaplan object from a csv file.
        """
        plan = Vaktaplan(df = pd.read_csv(file,index_col=0,header=0))
        return plan


    def get_date_range(self):
        """
        Helper function to get the first and last date of the shift plan.
        """
        first_day, first_month = map(int,self.df.columns[0].split('\n')[0].split('.'))
        last_day, last_month = map(int,self.df.columns[-1].split('\n')[0].split('.'))

        first_year = get_year(month=first_month)
        last_year = first_year if last_month >= first_month else first_year + 1

        first_date = datetime.date(year=int(first_year),month=int(first_month),day=int(first_day))
        last_date = datetime.date(year=int(last_year),month=int(last_month),day=int(last_day))
        return first_date,last_date



if __name__ == '__main__':
    main()